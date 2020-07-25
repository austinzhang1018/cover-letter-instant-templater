#! python3

import docx
from os.path import exists

def read_file(filename):
    try:
        doc = docx.Document(filename)
        # apply times new roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
    except:
        print('Invalid filename!')
        return None
    
    return doc

def get_template_fillables(filename):
    fields, stories = list(), list()

    template = read_file(filename)

    for paragraph in template.paragraphs:
        tagging = False
        fillable_name = ''

        for c in paragraph.text:
            if c == '{' or c == '[':
                if tagging:
                    raise Exception('Unexpected %s' %c)
                fillable_name = ''
                tagging = True
            elif c == '}' or c == ']':
                if not tagging:
                    raise Exception('Unexpected %s' %c)
                if c == '}':
                    fields.append(fillable_name)
                else:
                    stories.append(fillable_name)
                tagging = False
            elif tagging:
                fillable_name = ''.join((fillable_name, c))
    return fields, stories

def generate_form(company, template_filename):
    printed_fields = set()
    printed_stories = set()

    fields, stories = get_template_fillables(template_filename)

    width = max(len(max(fields, key=lambda field: len(field))), len(max(stories, key=lambda story: len(story)))) + 2

    with open('forms/%s.form' % company, 'w') as f:
        for field in fields:
            if field not in printed_fields:
                print(('%s:' % field).ljust(width), file=f)
                printed_fields.add(field)
        for story in stories:
            if story not in printed_stories:
                print(('Story | %s:' % story).ljust(width), file=f)
                printed_stories.add(story)
    print('form generated')
        
def get_form_info(filename):
    fields = dict()
    stories = dict()
    with open(filename, 'r') as f:
        for line in f:
            if 'Story |' in line:
                name, data = line.split('|')[1].strip().split(':')
                stories[name.strip()] = data.strip()
            else:
                name, data = line.split(':')
                fields[name.strip()] = data.strip()
    return fields, stories

def retreive_story(story_name):
    try:
        with open('stories/%s' % story_name) as f:
            return f.read()
    except:
        raise Exception('Story %s not found!' % story_name)

def replace_fillables(text, fillables, is_field=True):
    start, end = (text.find('{'), text.find('}')) if is_field else (text.find('['), text.find(']'))
    while start != -1 and end != -1:
        if end < start:
            raise Exception('First instance of } is before first instance of {. Make sure your template is defined correctly.')
        fillable = text[start+1:end]
        text = text.replace('{%s}' % fillable if is_field else '[%s]' % fillable, fillables[fillable])
        start, end = (text.find('{'), text.find('}')) if is_field else (text.find('['), text.find(']'))

    if start != -1 or end != -1:
        raise Exception('Unbalanced curly brackets. Make sure your template is defined correctly.')
    return text


def fill_template(company, template_filename):
    fields, stories = get_form_info('forms/%s.form' % company)
    for story in stories:
        stories[story] = retreive_story(stories[story])

    template = read_file(template_filename)

    for paragraph in template.paragraphs:        
        text = replace_fillables(paragraph.text, stories, is_field=False)
        text = replace_fillables(text, fields, is_field=True)
        if paragraph.text != text:
            paragraph.text = text
        
    template.save('letters/%s.docx' % company)


if __name__ == '__main__':
    TEMPLATE_NAME = 'cover-letter-template.docx'
    company = input('Please enter company name: ')
    action = input('Type form to create form. Type anything else to create a cover letter: ')
    if action == 'form':
        if not exists('forms/%s.form' % company) or 'continue' in input('A form for %s already exists. Type continue to overwrite and anything else to exit: ' % company):
            existing = input("If you'd like to copy an existing template, type the company name here. Otherwise, press enter: ")
            if existing != '':
                with open('forms/%s.form' % existing) as f:
                    with open('forms/%s.form' % company, "w") as f1:
                        for line in f:
                            f1.write(line)
            else:
                generate_form(company, TEMPLATE_NAME)
    else:
        fill_template(company, TEMPLATE_NAME)
